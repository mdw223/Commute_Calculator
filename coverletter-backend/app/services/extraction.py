"""Extract plain text from uploaded resume/profile documents."""

import io

import docx
from pypdf import PdfReader

SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}

MAX_EXTRACTED_CHARS = 40_000


class ExtractionError(Exception):
    pass


def extract_text(filename: str, content_type: str, data: bytes) -> str:
    lowered = filename.lower()
    try:
        if content_type == "application/pdf" or lowered.endswith(".pdf"):
            text = _extract_pdf(data)
        elif (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or lowered.endswith(".docx")
        ):
            text = _extract_docx(data)
        elif content_type == "text/plain" or lowered.endswith(".txt"):
            text = data.decode("utf-8", errors="ignore")
        else:
            raise ExtractionError(f"Unsupported file type: {content_type or lowered}")
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Failed to parse {filename}: {e}") from e

    text = text.strip()
    if not text:
        raise ExtractionError("No readable text found in document")
    return text[:MAX_EXTRACTED_CHARS]


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
