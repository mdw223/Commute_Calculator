"""One-off generator for the default cover letter .docx template.

Run this whenever the template layout needs to change:

    python scripts/generate_template.py

It writes app/templates/cover_letter_default.docx using python-docx, with
docxtpl (Jinja2-in-docx) placeholders for the fields produced by
app/services/docgen.py. Later, this file can simply be replaced with a
hand-designed template exported from Word/Google Docs as long as it keeps
the same placeholder names.
"""

import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_PATH = pathlib.Path(__file__).resolve().parent.parent / "app" / "templates" / "cover_letter_default.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x52, 0x52, 0x52)


def _set_base_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(8)


def build() -> None:
    document = Document()
    for section in document.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(65)
        section.right_margin = Pt(65)

    _set_base_style(document)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = name.add_run("{{ full_name }}")
    run.bold = True
    run.font.size = Pt(20)

    contact = document.add_paragraph()
    contact_run = contact.add_run(
        "{{ email }}{% if phone %} | {{ phone }}{% endif %}{% if location %} | {{ location }}{% endif %}"
    )
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = MUTED
    contact.paragraph_format.space_after = Pt(20)

    date_p = document.add_paragraph("{{ date }}")
    date_p.paragraph_format.space_after = Pt(16)

    recipient = document.add_paragraph()
    recipient.add_run("{{ hiring_manager }}")
    recipient_company = document.add_paragraph()
    recipient_company.paragraph_format.space_after = Pt(2)
    recipient_company.add_run("{{ company_name }}")
    recipient_company.paragraph_format.space_after = Pt(20)

    greeting = document.add_paragraph("Dear {{ hiring_manager }},")
    greeting.paragraph_format.space_after = Pt(12)

    opening = document.add_paragraph("{{ opening_paragraph }}")
    opening.paragraph_format.space_after = Pt(12)

    # Dynamic number of body paragraphs — docxtpl's `{%p %}` paragraph tag
    # removes the two "loop marker" paragraphs and repeats the middle one
    # once per item in body_paragraphs.
    loop_start = document.add_paragraph("{%p for para in body_paragraphs %}")
    body = document.add_paragraph("{{ para }}")
    body.paragraph_format.space_after = Pt(12)
    loop_end = document.add_paragraph("{%p endfor %}")
    for p in (loop_start, loop_end):
        p.paragraph_format.space_after = Pt(0)

    closing = document.add_paragraph("{{ closing_paragraph }}")
    closing.paragraph_format.space_after = Pt(24)

    sign_off = document.add_paragraph("{{ sign_off }},")
    sign_off.paragraph_format.space_after = Pt(36)

    signature = document.add_paragraph("{{ full_name }}")
    signature.add_run().bold = True

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Wrote template to {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
