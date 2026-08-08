import io

import docx

from app.services.docgen import build_filename, render_cover_letter_docx


def test_render_cover_letter_docx_includes_all_fields():
    content = {
        "company_name": "Acme Corp",
        "job_title": "Software Engineer",
        "hiring_manager": "Alex Smith",
        "opening_paragraph": "I'm excited to apply for the Software Engineer role.",
        "body_paragraphs": [
            "I built several key features at my last job.",
            "I admire Acme's mission and would love to contribute.",
        ],
        "closing_paragraph": "Thank you for your consideration.",
        "sign_off": "Sincerely",
    }

    data = render_cover_letter_docx(
        content=content,
        full_name="Jane Doe",
        email="jane@example.com",
        phone="555-123-4567",
        location="Austin, TX",
    )

    document = docx.Document(io.BytesIO(data))
    full_text = "\n".join(p.text for p in document.paragraphs)

    assert "Jane Doe" in full_text
    assert "jane@example.com" in full_text
    assert "555-123-4567" in full_text
    assert "Austin, TX" in full_text
    assert "Acme Corp" in full_text
    assert "Dear Alex Smith," in full_text
    assert content["opening_paragraph"] in full_text
    for para in content["body_paragraphs"]:
        assert para in full_text
    assert content["closing_paragraph"] in full_text
    assert "Sincerely," in full_text


def test_render_defaults_hiring_manager_when_missing():
    content = {
        "company_name": "Acme Corp",
        "job_title": "Engineer",
        "opening_paragraph": "Opening.",
        "body_paragraphs": ["Body."],
        "closing_paragraph": "Closing.",
    }
    data = render_cover_letter_docx(
        content=content, full_name="Jane Doe", email="jane@example.com", phone=None, location=None
    )
    document = docx.Document(io.BytesIO(data))
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "Dear Hiring Manager," in full_text
    assert "Sincerely," in full_text  # default sign-off


def test_build_filename_sanitizes_and_falls_back():
    assert build_filename("Acme/Corp:", "Sr. Engineer?") == "Cover Letter - Sr. Engineer - AcmeCorp.docx"
    assert build_filename(None, None) == "Cover Letter - Cover Letter.docx"
