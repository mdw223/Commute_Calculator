import io

import docx
import pytest

from app.services.extraction import ExtractionError, extract_text


def test_extract_plain_text():
    text = extract_text("notes.txt", "text/plain", b"Hello, this is my resume text.")
    assert text == "Hello, this is my resume text."


def test_extract_docx():
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Software Engineer with 5 years of experience.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text(
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )
    assert "Jane Doe" in text
    assert "Software Engineer" in text


def test_extract_rejects_unsupported_type():
    with pytest.raises(ExtractionError):
        extract_text("image.png", "image/png", b"\x89PNG\r\n")


def test_extract_rejects_empty_document():
    document = docx.Document()
    buffer = io.BytesIO()
    document.save(buffer)
    with pytest.raises(ExtractionError):
        extract_text(
            "empty.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buffer.getvalue(),
        )


def test_extract_truncates_long_text():
    long_text = "word " * 20000
    text = extract_text("long.txt", "text/plain", long_text.encode())
    from app.services.extraction import MAX_EXTRACTED_CHARS

    assert len(text) <= MAX_EXTRACTED_CHARS
